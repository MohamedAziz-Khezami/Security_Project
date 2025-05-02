import streamlit as st
from cryptography.hazmat.primitives import hashes, serialization
from cryptography.hazmat.primitives.asymmetric import rsa, padding as asym_padding
from cryptography.hazmat.primitives.ciphers import Cipher, algorithms, modes
from cryptography.hazmat.primitives.ciphers.aead import AESGCM, ChaCha20Poly1305
# Import Camellia
from cryptography.hazmat.primitives.ciphers.algorithms import Camellia
from cryptography.hazmat.primitives.kdf.pbkdf2 import PBKDF2HMAC
# Import Scrypt
from cryptography.hazmat.primitives.kdf.scrypt import Scrypt
from cryptography.hazmat.backends import default_backend
from cryptography.exceptions import InvalidTag, InvalidSignature # For specific error handling
import os
import io
import base64
from docx import Document
import tempfile
import re # For partial decryption pattern matching

# --- Configuration Constants ---
# Symmetric Algo Settings
SYMMETRIC_KEY_LENGTH = 32 # 32 bytes = 256 bits for AES/Camellia/ChaCha20
SALT_LENGTH = 16
AES_IV_LENGTH = 12 # Recommended IV length for AES-GCM
AES_TAG_LENGTH = 16 # Standard tag length for AES-GCM
CAMELLIA_IV_LENGTH = 12 # GCM standard
CAMELLIA_TAG_LENGTH = 16
CHACHA20_NONCE_LENGTH = 12 # Standard nonce length for ChaCha20Poly1305
# KDF Default Settings
DEFAULT_PBKDF2_ITERATIONS = 150000 # Slightly increased default
DEFAULT_SCRYPT_N = 2**14 # CPU/memory cost factor (power of 2)
DEFAULT_SCRYPT_R = 8     # Block size
DEFAULT_SCRYPT_P = 1     # Parallelization factor
# RSA Settings
RSA_KEY_SIZE = 2048
RSA_PUBLIC_EXPONENT = 65537
RSA_OAEP_PADDING = asym_padding.OAEP(
    mgf=asym_padding.MGF1(algorithm=hashes.SHA256()),
    algorithm=hashes.SHA256(),
    label=None
)
# Partial Encryption Marker
PARTIAL_ENC_MARKER_START = "::ENC["
PARTIAL_ENC_MARKER_END = "]::"
PARTIAL_ENC_PATTERN = re.compile(r"::ENC\[([A-Za-z0-9+/=]+)\]::")

# --- Setup ---
st.set_page_config(page_title="SecureCrypt Pro", layout="wide", initial_sidebar_state="expanded")
st.title("🔐 SecureCrypt Pro: Advanced File Encryption & Decryption")
st.markdown("Upload a file, choose an operation and algorithm, provide necessary credentials, and download the result.")
backend = default_backend()

# --- Cryptographic Functions ---

# Key Derivation (Now supports PBKDF2 and Scrypt)
def derive_key(password: str, salt: bytes, kdf_type: str, **kdf_params) -> bytes:
    """Derives a cryptographic key from a password using the specified KDF."""
    password_bytes = password.encode()

    if kdf_type == 'PBKDF2':
        iterations = kdf_params.get('iterations', DEFAULT_PBKDF2_ITERATIONS)
        st.write(f"_(Using PBKDF2 with {iterations} iterations)_") # Debug/Info
        kdf = PBKDF2HMAC(
            algorithm=hashes.SHA256(),
            length=SYMMETRIC_KEY_LENGTH,
            salt=salt,
            iterations=iterations,
            backend=backend
        )
        return kdf.derive(password_bytes)
    elif kdf_type == 'Scrypt':
        n = kdf_params.get('n', DEFAULT_SCRYPT_N)
        r = kdf_params.get('r', DEFAULT_SCRYPT_R)
        p = kdf_params.get('p', DEFAULT_SCRYPT_P)
        st.write(f"_(Using Scrypt with n={n}, r={r}, p={p})_") # Debug/Info
        kdf = Scrypt(
            salt=salt,
            length=SYMMETRIC_KEY_LENGTH,
            n=n,
            r=r,
            p=p,
            backend=backend
        )
        return kdf.derive(password_bytes)
    else:
        raise ValueError("Unsupported KDF type specified.")

# --- AES-GCM --- (Unchanged from previous version, but now uses generic derive_key)
def aes_gcm_encrypt(data: bytes, password: str, kdf_type: str, **kdf_params) -> bytes:
    salt = os.urandom(SALT_LENGTH)
    key = derive_key(password, salt, kdf_type, **kdf_params)
    iv = os.urandom(AES_IV_LENGTH)
    aesgcm = AESGCM(key)
    encrypted = aesgcm.encrypt(iv, data, None)
    return salt + iv + encrypted

def aes_gcm_decrypt(encrypted_data: bytes, password: str, kdf_type: str, **kdf_params) -> bytes:
    try:
        salt = encrypted_data[:SALT_LENGTH]
        iv = encrypted_data[SALT_LENGTH : SALT_LENGTH + AES_IV_LENGTH]
        ciphertext_with_tag = encrypted_data[SALT_LENGTH + AES_IV_LENGTH:]
        key = derive_key(password, salt, kdf_type, **kdf_params)
        aesgcm = AESGCM(key)
        return aesgcm.decrypt(iv, ciphertext_with_tag, None)
    except InvalidTag:
        st.error("AES Decryption Failed: Invalid Tag. Password may be incorrect or data corrupted.")
        raise
    except Exception as e:
        st.error(f"AES Decryption Failed: {e}")
        raise

# --- ChaCha20Poly1305 --- (Unchanged from previous version, but now uses generic derive_key)
def chacha20_encrypt(data: bytes, password: str, kdf_type: str, **kdf_params) -> bytes:
    salt = os.urandom(SALT_LENGTH)
    key = derive_key(password, salt, kdf_type, **kdf_params)
    nonce = os.urandom(CHACHA20_NONCE_LENGTH)
    chacha = ChaCha20Poly1305(key)
    encrypted = chacha.encrypt(nonce, data, None)
    return salt + nonce + encrypted

def chacha20_decrypt(encrypted_data: bytes, password: str, kdf_type: str, **kdf_params) -> bytes:
    try:
        salt = encrypted_data[:SALT_LENGTH]
        nonce = encrypted_data[SALT_LENGTH : SALT_LENGTH + CHACHA20_NONCE_LENGTH]
        ciphertext_with_tag = encrypted_data[SALT_LENGTH + CHACHA20_NONCE_LENGTH:]
        key = derive_key(password, salt, kdf_type, **kdf_params)
        chacha = ChaCha20Poly1305(key)
        return chacha.decrypt(nonce, ciphertext_with_tag, None)
    except InvalidTag:
        st.error("ChaCha20 Decryption Failed: Invalid Tag. Password may be incorrect or data corrupted.")
        raise
    except Exception as e:
        st.error(f"ChaCha20 Decryption Failed: {e}")
        raise

# --- Camellia-GCM --- (NEW)
def camellia_gcm_encrypt(data: bytes, password: str, kdf_type: str, **kdf_params) -> bytes:
    """Encrypts data using Camellia-GCM with a password-derived key."""
    salt = os.urandom(SALT_LENGTH)
    key = derive_key(password, salt, kdf_type, **kdf_params)
    iv = os.urandom(CAMELLIA_IV_LENGTH)
    # Use Cipher and GCM mode for Camellia AEAD
    cipher = Cipher(Camellia(key), modes.GCM(iv), backend=backend)
    encryptor = cipher.encryptor()
    # encryptor.authenticate_additional_data(b"") # No associated data
    encrypted = encryptor.update(data) + encryptor.finalize()
    # Store as: salt + iv + ciphertext + tag (tag is appended by GCM mode)
    return salt + iv + encrypted + encryptor.tag # Manually append tag

def camellia_gcm_decrypt(encrypted_data: bytes, password: str, kdf_type: str, **kdf_params) -> bytes:
    """Decrypts data encrypted with Camellia-GCM."""
    try:
        salt = encrypted_data[:SALT_LENGTH]
        iv = encrypted_data[SALT_LENGTH : SALT_LENGTH + CAMELLIA_IV_LENGTH]
        # Extract tag (last CAMELLIA_TAG_LENGTH bytes)
        tag = encrypted_data[-CAMELLIA_TAG_LENGTH:]
        # Ciphertext is between iv and tag
        ciphertext = encrypted_data[SALT_LENGTH + CAMELLIA_IV_LENGTH : -CAMELLIA_TAG_LENGTH]
        key = derive_key(password, salt, kdf_type, **kdf_params)
        # Use Cipher and GCM mode for Camellia AEAD
        cipher = Cipher(Camellia(key), modes.GCM(iv, tag), backend=backend)
        decryptor = cipher.decryptor()
        # decryptor.authenticate_additional_data(b"") # No associated data
        return decryptor.update(ciphertext) + decryptor.finalize()
    except InvalidTag:
        st.error("Camellia Decryption Failed: Invalid Tag. Password may be incorrect or data corrupted.")
        raise
    except Exception as e:
        st.error(f"Camellia Decryption Failed: {e}")
        raise

# --- RSA Key Handling --- (Unchanged structure, added password handling)
def generate_rsa_keys():
    private_key = rsa.generate_private_key(
        public_exponent=RSA_PUBLIC_EXPONENT,
        key_size=RSA_KEY_SIZE,
        backend=backend
    )
    public_key = private_key.public_key()
    return private_key, public_key

def serialize_private_key(private_key, password: str = None) -> bytes:
    if password:
        enc_alg = serialization.BestAvailableEncryption(password.encode())
    else:
        enc_alg = serialization.NoEncryption()
    pem = private_key.private_bytes(
        encoding=serialization.Encoding.PEM,
        format=serialization.PrivateFormat.PKCS8,
        encryption_algorithm=enc_alg
    )
    return pem

def serialize_public_key(public_key) -> bytes:
    pem = public_key.public_bytes(
        encoding=serialization.Encoding.PEM,
        format=serialization.PublicFormat.SubjectPublicKeyInfo
    )
    return pem

def load_private_key(pem_data: bytes, password: str = None):
    try:
        password_bytes = password.encode() if password else None
        private_key = serialization.load_pem_private_key(
            pem_data,
            password=password_bytes,
            backend=backend
        )
        return private_key
    except (ValueError, TypeError) as e: # Catches incorrect password or format errors
        if "Bad decrypt" in str(e) or "MAC check failed" in str(e):
             st.error("Failed to load private key: Incorrect password.")
        else:
             st.error(f"Failed to load private key: Invalid key format or corrupted file. {e}")
        return None
    except Exception as e:
        st.error(f"An unexpected error occurred loading the private key: {e}")
        return None

# --- RSA Hybrid Encryption --- (Unchanged logic)
def rsa_hybrid_encrypt(data: bytes, public_key) -> bytes:
    symmetric_key = AESGCM.generate_key(bit_length=256)
    aesgcm = AESGCM(symmetric_key)
    iv = os.urandom(AES_IV_LENGTH)
    encrypted_data = aesgcm.encrypt(iv, data, None)
    encrypted_symmetric_key = public_key.encrypt(symmetric_key, RSA_OAEP_PADDING)
    encrypted_key_len = len(encrypted_symmetric_key).to_bytes(2, 'big')
    return encrypted_key_len + encrypted_symmetric_key + iv + encrypted_data

def rsa_hybrid_decrypt(encrypted_payload: bytes, private_key) -> bytes:
    try:
        encrypted_key_len = int.from_bytes(encrypted_payload[:2], 'big')
        offset = 2
        encrypted_symmetric_key = encrypted_payload[offset : offset + encrypted_key_len]
        offset += encrypted_key_len
        iv = encrypted_payload[offset : offset + AES_IV_LENGTH]
        offset += AES_IV_LENGTH
        encrypted_data = encrypted_payload[offset:]
        symmetric_key = private_key.decrypt(encrypted_symmetric_key, RSA_OAEP_PADDING)
        aesgcm = AESGCM(symmetric_key)
        return aesgcm.decrypt(iv, encrypted_data, None)
    except InvalidTag:
        st.error("RSA Hybrid Decryption Failed: AES-GCM integrity check failed (Invalid Tag). Data may be corrupted.")
        raise
    except (InvalidSignature, ValueError) as e: # Catch RSA decryption errors
         st.error(f"RSA Hybrid Decryption Failed: RSA key decryption failed. Incorrect private key? {e}")
         raise
    except Exception as e:
        st.error(f"RSA Hybrid Decryption Failed: {e}")
        raise

# --- Helper Functions ---
def get_file_extension(filename):
    return os.path.splitext(filename)[1].lower()

def get_mime_type(filename):
    ext = get_file_extension(filename)
    # Add more mappings as needed
    mime_map = {
        ".txt": "text/plain", ".pdf": "application/pdf", ".jpg": "image/jpeg",
        ".jpeg": "image/jpeg", ".png": "image/png",
        ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ".md": "text/markdown", ".csv": "text/csv", ".json": "application/json",
        ".zip": "application/zip", ".bin": "application/octet-stream",
        ".pem": "application/x-pem-file"
        }
    return mime_map.get(ext, "application/octet-stream") # Default binary


# --- Streamlit UI ---

# Sidebar Layout
st.sidebar.header("⚙️ 1. Select File & Operation")
uploaded_file = st.sidebar.file_uploader(
    "Upload a file",
    label_visibility="collapsed",
    type=None # Allow any file type initially
)
operation = st.sidebar.radio("Operation", ["Encrypt", "Decrypt"], horizontal=True, key="op_radio")

st.sidebar.header("⚙️ 2. Choose Algorithm")

# Algorithm descriptions (Tooltips)
algo_tooltips = {
    "AES-GCM (Symmetric)": "Advanced Encryption Standard (256-bit) with Galois/Counter Mode. Provides confidentiality and integrity. Widely used standard.",
    "ChaCha20Poly1305 (Symmetric)": "Stream cipher combined with Poly1305 authenticator. Fast on software, excellent security. Good alternative to AES-GCM.",
    "Camellia-GCM (Symmetric)": "Block cipher (256-bit) similar to AES, also using GCM for authenticated encryption. Less common than AES but secure.",
    "RSA (Asymmetric - Hybrid)": "Uses RSA (2048-bit) to encrypt a temporary symmetric key (AES-GCM), which encrypts the data. Standard for asymmetric encryption of files. Requires key pair."
}
algo_options = list(algo_tooltips.keys())
algorithm = st.sidebar.selectbox(
    "Encryption Algorithm",
    algo_options,
    help=algo_tooltips[algo_options[0]], # Initial help text
    key="algo_select"
)
# Update help text dynamically (minor UX improvement)
st.sidebar.markdown(f"<small>{algo_tooltips[algorithm]}</small>", unsafe_allow_html=True)


st.sidebar.header("⚙️ 3. Configure Settings")

# --- Algorithm Specific UI Elements ---
password = None
kdf_type = None
kdf_params = {}
private_key_password = None
private_key_file = None
partial = False
partial_text_scope = "First Occurrence" # Default scope

# Symmetric Algorithm Settings
if "Symmetric" in algorithm:
    st.sidebar.subheader(f"{algorithm.split(' ')[0]} Options")
    password = st.sidebar.text_input("🔑 Password", type="password", key="sym_password")

    # KDF Selection
    kdf_type = st.sidebar.radio(
        "Key Derivation Function (KDF)",
        ['PBKDF2', 'Scrypt'],
        horizontal=True,
        help="Method to turn your password into an encryption key. Scrypt is generally stronger against custom hardware attacks but slower. **Must match KDF & parameters used for encryption/decryption!**"
    )

    # Advanced KDF Parameters
    with st.sidebar.expander("Advanced KDF Settings"):
        st.warning("**Important:** You must use the **exact same KDF parameters** for decryption as were used during encryption.", icon="⚠️")
        if kdf_type == 'PBKDF2':
            pbkdf2_iterations = st.number_input("PBKDF2 Iterations", min_value=10000, value=DEFAULT_PBKDF2_ITERATIONS, step=10000, help="Higher = more secure, but slower.")
            kdf_params['iterations'] = pbkdf2_iterations
        elif kdf_type == 'Scrypt':
            scrypt_n = st.number_input("Scrypt N (Cost)", min_value=2, value=DEFAULT_SCRYPT_N, step=1024, help="CPU/memory cost factor (power of 2). Higher = more secure, slower, more memory.")
            scrypt_r = st.number_input("Scrypt R (Block Size)", min_value=1, value=DEFAULT_SCRYPT_R, step=1, help="Block size parameter.")
            scrypt_p = st.number_input("Scrypt P (Parallelism)", min_value=1, value=DEFAULT_SCRYPT_P, step=1, help="Parallelization factor.")
            kdf_params['n'] = scrypt_n
            kdf_params['r'] = scrypt_r
            kdf_params['p'] = scrypt_p

    # Partial Encryption Checkbox
    if uploaded_file and uploaded_file.type in ["text/plain", "application/vnd.openxmlformats-officedocument.wordprocessingml.document", "text/markdown", "text/csv", "application/json"]:
         partial = st.sidebar.checkbox(f"Partial {operation} (Text-based files)", help="Encrypt/Decrypt specific text snippets within the file instead of the whole file.")
         if partial and operation == "Encrypt":
             partial_text_scope = st.sidebar.radio("Encrypt Scope", ["First Occurrence", "All Occurrences"], index=0, horizontal=True)


# Asymmetric Algorithm Settings
elif "RSA" in algorithm:
    st.sidebar.subheader("RSA Options")
    if operation == "Encrypt":
        st.sidebar.info("A new RSA key pair (2048-bit) will be generated for encryption.")
        private_key_password = st.sidebar.text_input(
            "🔒 Password for new Private Key (Optional)",
            type="password",
            key="rsa_gen_pwd",
            help="Encrypts the downloaded private key file. Leave blank for no password."
        )
    else: # Decrypt
        private_key_file = st.sidebar.file_uploader(
            "🔑 Upload Private Key (.pem)",
            type=["pem"],
            key="rsa_priv_key_upload",
            help="Upload the `.pem` private key file corresponding to the public key used for encryption."
        )
        private_key_password = st.sidebar.text_input(
            "🔒 Password for Private Key (if set)",
            type="password",
            key="rsa_load_pwd",
            help="Enter the password if the private key file was encrypted."
        )
        if not private_key_file and uploaded_file:
            st.sidebar.warning("Private key file is required for RSA decryption.")


# --- Main Processing Area ---

def handle_symmetric_operation(file_bytes, file_name, file_type):
    """Handles symmetric encryption and decryption logic."""
    if not password:
        st.warning("⚠️ Please enter a password in the sidebar.")
        return None, None, None

    output_data = None
    output_filename = f"{operation.lower()}ed_{file_name}"
    output_mime = get_mime_type(output_filename.replace(f"{operation.lower()}ed_", "")) # Guess original type

    # Select appropriate functions based on algorithm
    algo_funcs = {
        "AES-GCM (Symmetric)": (aes_gcm_encrypt, aes_gcm_decrypt),
        "ChaCha20Poly1305 (Symmetric)": (chacha20_encrypt, chacha20_decrypt),
        "Camellia-GCM (Symmetric)": (camellia_gcm_encrypt, camellia_gcm_decrypt)
    }
    encrypt_func, decrypt_func = algo_funcs[algorithm]

    status_text = f"{operation}ing using {algorithm.split(' ')[0]} with {kdf_type} KDF..."
    with st.status(status_text, expanded=True) as status:
        try:
            st.write("Reading file...")
            # --- Partial Operation ---
            if partial:
                st.write("Performing partial operation...")
                # Supported types for partial operations
                supported_partial_types = ["text/plain", "text/markdown", "text/csv", "application/json",
                                           "application/vnd.openxmlformats-officedocument.wordprocessingml.document"]

                if file_type not in supported_partial_types:
                    status.error("Partial operation is only supported for text-based files (txt, md, csv, json, docx).")
                    return None, None, None

                if file_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                    # --- Partial DOCX ---
                    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
                        tmp.write(file_bytes)
                        tmp_path = tmp.name
                    try:
                        document = Document(tmp_path)
                        processed = False
                        if operation == "Encrypt":
                            st.subheader("Partial Encryption (DOCX)")
                            col1, col2 = st.columns(2)
                            with col1:
                                full_text_preview = "\n".join([p.text for p in document.paragraphs[:20]]) # Preview first 20 paragraphs
                                st.text_area("Original Content Preview (Paragraphs)", full_text_preview, height=200, key="docx_orig_preview", disabled=True)
                            with col2:
                                to_encrypt = st.text_area("Enter the exact text snippet to encrypt:", height=100, key="partial_enc_docx")
                                if to_encrypt:
                                    replacements_made = 0
                                    for paragraph in document.paragraphs:
                                        if to_encrypt in paragraph.text:
                                            encrypted_part_bytes = encrypt_func(to_encrypt.encode(), password, kdf_type, **kdf_params)
                                            encrypted_part_b64 = base64.b64encode(encrypted_part_bytes).decode('utf-8')
                                            replace_count = 1 if partial_text_scope == "First Occurrence" else -1 # -1 means all
                                            new_text = paragraph.text.replace(to_encrypt, f"{PARTIAL_ENC_MARKER_START}{encrypted_part_b64}{PARTIAL_ENC_MARKER_END}", replace_count)
                                            if new_text != paragraph.text:
                                                paragraph.text = new_text
                                                replacements_made += paragraph.text.count(f"{PARTIAL_ENC_MARKER_START}{encrypted_part_b64}{PARTIAL_ENC_MARKER_END}") # Rough count
                                                processed = True
                                                if partial_text_scope == "First Occurrence":
                                                    break # Only replace first found instance overall
                                    if processed:
                                        status.update(label=f"Encrypted {replacements_made} occurrence(s). Saving DOCX...")
                                        output_io = io.BytesIO()
                                        document.save(output_io)
                                        output_io.seek(0)
                                        output_data = output_io.read() # Read bytes from BytesIO
                                        output_filename = f"partially_encrypted_{file_name}"
                                        output_mime = file_type
                                        st.success(f"Snippet found and encrypted in {replacements_made} place(s).")
                                    else:
                                        st.warning("The text snippet to encrypt was not found in the DOCX paragraphs.")
                                else:
                                    st.info("Enter the text snippet to encrypt in the box above.")

                        else: # Decrypt DOCX
                             st.subheader("Partial Decryption (DOCX)")
                             decryptions_made = 0
                             for paragraph in document.paragraphs:
                                 original_text = paragraph.text
                                 new_text = original_text
                                 # Find all encrypted blocks in the paragraph
                                 for match in PARTIAL_ENC_PATTERN.finditer(original_text):
                                     encrypted_part_b64 = match.group(1)
                                     block_to_replace = match.group(0) # Full ::ENC[...]:: block
                                     st.write(f"Found potential block: {block_to_replace[:20]}...")
                                     try:
                                         encrypted_part_bytes = base64.b64decode(encrypted_part_b64)
                                         decrypted_part = decrypt_func(encrypted_part_bytes, password, kdf_type, **kdf_params).decode()
                                         # Replace only this specific matched block
                                         new_text = new_text.replace(block_to_replace, decrypted_part, 1)
                                         decryptions_made += 1
                                         processed = True
                                         st.write("Decryption successful for this block.")
                                     except Exception as e:
                                         st.warning(f"Could not decrypt block {block_to_replace[:20]}... Skipping. Error: {e}")
                                 if new_text != original_text:
                                     paragraph.text = new_text

                             if processed:
                                status.update(label=f"Decrypted {decryptions_made} block(s). Saving DOCX...")
                                output_io = io.BytesIO()
                                document.save(output_io)
                                output_io.seek(0)
                                output_data = output_io.read()
                                output_filename = f"partially_decrypted_{file_name}"
                                output_mime = file_type
                                st.success(f"Found and decrypted {decryptions_made} block(s) in the DOCX.")
                             else:
                                st.warning(f"Could not find any encrypted blocks ({PARTIAL_ENC_MARKER_START}...{PARTIAL_ENC_MARKER_END}) in the DOCX paragraphs.")

                    finally:
                        os.unlink(tmp_path) # Clean up temp file

                else:
                    # --- Partial Text/CSV/JSON/MD ---
                    try:
                        original_text = file_bytes.decode() # Assume text-based
                    except UnicodeDecodeError:
                        status.error("Failed to decode file as text. Ensure it's a valid text-based file (txt, md, csv, json).")
                        return None, None, None

                    if operation == "Encrypt":
                        st.subheader("Partial Encryption")
                        col1, col2 = st.columns(2)
                        with col1:
                            st.text_area("Original Content", original_text, height=300, key="text_orig", disabled=True)
                        with col2:
                            to_encrypt = st.text_area("Enter the exact text snippet to encrypt:", height=100, key="partial_enc_text")
                            if to_encrypt and to_encrypt in original_text:
                                st.write("Encrypting snippet...")
                                encrypted_part_bytes = encrypt_func(to_encrypt.encode(), password, kdf_type, **kdf_params)
                                encrypted_part_b64 = base64.b64encode(encrypted_part_bytes).decode('utf-8')
                                replace_count = 1 if partial_text_scope == "First Occurrence" else -1
                                partially_encrypted_text = original_text.replace(to_encrypt, f"{PARTIAL_ENC_MARKER_START}{encrypted_part_b64}{PARTIAL_ENC_MARKER_END}", replace_count)
                                replacements_made = original_text.count(to_encrypt) if replace_count == -1 else (1 if to_encrypt in original_text else 0)
                                st.text_area("Partially Encrypted Content", partially_encrypted_text, height=150, key="text_enc_result")
                                output_data = partially_encrypted_text.encode()
                                output_filename = f"partially_encrypted_{file_name}"
                                output_mime = file_type
                                status.update(label=f"Encrypted {replacements_made} occurrence(s).")
                                processed = True
                            elif to_encrypt:
                                st.warning("The text snippet to encrypt was not found in the original content.")
                            else:
                                st.info("Enter the text snippet you wish to encrypt above.")

                    else: # Decrypt Text
                        st.subheader("Partial Decryption")
                        col1, col2 = st.columns(2)
                        with col1:
                             st.text_area("Encrypted Content", original_text, height=300, key="text_enc", disabled=True)

                        decrypted_text = original_text
                        decryptions_made = 0
                        matches = list(PARTIAL_ENC_PATTERN.finditer(original_text)) # Find all blocks first

                        if not matches:
                             st.warning(f"Could not find any encrypted blocks ({PARTIAL_ENC_MARKER_START}...{PARTIAL_ENC_MARKER_END}).")
                        else:
                            st.write(f"Found {len(matches)} potential encrypted block(s). Attempting decryption...")
                            processed_text = original_text
                            for match in matches:
                                encrypted_part_b64 = match.group(1)
                                block_to_replace = match.group(0)
                                st.write(f"- Found block: {block_to_replace[:20]}...")
                                try:
                                    encrypted_part_bytes = base64.b64decode(encrypted_part_b64)
                                    decrypted_part = decrypt_func(encrypted_part_bytes, password, kdf_type, **kdf_params).decode()
                                    # Replace only the first remaining occurrence of this specific block
                                    processed_text = processed_text.replace(block_to_replace, decrypted_part, 1)
                                    decryptions_made += 1
                                    processed = True
                                    st.write("  ✅ Decryption successful.")
                                except Exception as e:
                                    st.warning(f"  ❌ Could not decrypt block {block_to_replace[:20]}... Skipping. Incorrect password/params or invalid block? Error: {e}")

                            if processed:
                                with col2:
                                    st.text_area("Partially Decrypted Content", processed_text, height=300, key="text_dec_result")
                                output_data = processed_text.encode()
                                output_filename = f"partially_decrypted_{file_name}"
                                output_mime = file_type
                                status.update(label=f"Decrypted {decryptions_made} block(s).")
                            else:
                                # Display original text if no successful decryptions
                                with col2:
                                    st.text_area("Decryption Result", original_text, height=300, key="text_dec_fail", disabled=True)


            # --- Full File Operation ---
            else:
                st.write("Performing full file operation...")
                if operation == "Encrypt":
                    output_data = encrypt_func(file_bytes, password, kdf_type, **kdf_params)
                    output_mime = "application/octet-stream" # Encrypted data is binary
                    status.update(label=f"File encrypted successfully using {algorithm.split(' ')[0]}!", state="complete")
                else: # Decrypt
                    output_data = decrypt_func(file_bytes, password, kdf_type, **kdf_params)
                    # output_mime already guessed above
                    status.update(label=f"File decrypted successfully using {algorithm.split(' ')[0]}!", state="complete")

            return output_data, output_filename, output_mime

        except Exception as e:
            status.error(f"Operation failed: {e}")
            # import traceback
            # st.error(traceback.format_exc()) # Uncomment for detailed debug logs
            return None, None, None


def handle_asymmetric_operation(file_bytes, file_name):
    """Handles asymmetric RSA hybrid encryption and decryption."""
    output_data = None
    output_filename = f"{operation.lower()}ed_{file_name}"
    output_mime = "application/octet-stream" # Encrypted is binary, decrypted might be binary

    status_text = f"{operation}ing using {algorithm}..."
    with st.status(status_text, expanded=True) as status:
        try:
            if operation == "Encrypt":
                st.write("Generating RSA key pair...")
                private_key, public_key = generate_rsa_keys()
                st.write(f"Encrypting file using temporary AES key protected by RSA public key...")
                output_data = rsa_hybrid_encrypt(file_bytes, public_key)

                # Serialize keys for download
                st.write("Serializing keys for download...")
                private_pem = serialize_private_key(private_key, private_key_password if private_key_password else None)
                public_pem = serialize_public_key(public_key)

                status.update(label="File encrypted successfully! Download keys below.", state="complete")

                # Display download buttons for keys in an expander
                with st.expander("🔑 Download Generated Keys (Required for Decryption!)", expanded=True):
                     col1, col2 = st.columns(2)
                     with col1:
                         st.download_button(
                             label="Download Private Key (.pem)",
                             data=private_pem,
                             file_name="private_key.pem",
                             mime="application/x-pem-file",
                             help="KEEP THIS SAFE! Needed to decrypt the file. Store securely."
                         )
                     with col2:
                         st.download_button(
                             label="Download Public Key (.pem)",
                             data=public_pem,
                             file_name="public_key.pem",
                             mime="application/x-pem-file",
                             help="Can be shared to allow others to encrypt files for you."
                         )
                     if private_key_password:
                         st.info("ℹ️ The private key is password protected.")
                     else:
                         st.warning("⚠️ The private key is NOT password protected. Handle with extra care!")

            else: # Decrypt
                if not private_key_file:
                    st.warning("⚠️ Please upload your private key (.pem file) in the sidebar.")
                    status.update(label="Waiting for private key upload.", state="error")
                    return None, None, None

                st.write("Loading private key...")
                private_key_pem = private_key_file.read()
                private_key = load_private_key(private_key_pem, private_key_password if private_key_password else None)

                if not private_key:
                    # Error message already shown by load_private_key
                    status.update(label="Failed to load private key.", state="error")
                    return None, None, None

                st.write("Private key loaded. Decrypting file...")
                output_data = rsa_hybrid_decrypt(file_bytes, private_key)
                # Try to guess original mime type based on the original filename
                output_mime = get_mime_type(file_name.replace("encrypted_", "", 1))
                status.update(label="File decrypted successfully!", state="complete")

            return output_data, output_filename, output_mime

        except Exception as e:
            status.error(f"Operation failed: {e}")
            # import traceback
            # st.error(traceback.format_exc()) # Uncomment for detailed debug logs
            return None, None, None

# --- Main Execution Logic ---
if uploaded_file:
    st.divider()
    st.header(f"🚀 Processing: {operation} '{uploaded_file.name}'")

    file_bytes_content = uploaded_file.read()
    file_name_content = uploaded_file.name
    file_type_content = uploaded_file.type

    processed_data = None
    processed_filename = None
    processed_mime = None

    if "Symmetric" in algorithm:
        processed_data, processed_filename, processed_mime = handle_symmetric_operation(
            file_bytes_content, file_name_content, file_type_content
        )
    elif "RSA" in algorithm:
         processed_data, processed_filename, processed_mime = handle_asymmetric_operation(
            file_bytes_content, file_name_content
        )

    # Offer download if data was processed successfully
    if processed_data is not None and processed_filename is not None:
        st.divider()
        st.header("✅ Results & Download")
        st.success(f"Operation completed successfully!")
        st.download_button(
            label=f"💾 Download {processed_filename}",
            data=processed_data,
            file_name=processed_filename,
            mime=processed_mime,
            key="download_button"
        )

else:
    st.info("☝️ Upload a file using the sidebar to begin.")


# Footer info
st.sidebar.markdown("---")
st.sidebar.info("SecureCrypt Pro v2.0\n\nDeveloped using Streamlit & Python Cryptography library.")